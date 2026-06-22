#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成 ZUUL — 失落的古迹 小组实训报告 (REPORT.docx)
严格按六级大纲结构编写，所有数据来自实际项目代码与开发记录。
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import datetime

doc = Document()

# ========== 全局样式 ==========
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


def set_cell_shading(cell, color_str):
    shading_elm = cell._tc.get_or_add_tcPr()
    shading = shading_elm.makeelement(qn('w:shd'), {qn('w:fill'): color_str, qn('w:val'): 'clear'})
    shading_elm.append(shading)


def make_header_row(table, row_idx, texts, color='4472C4'):
    row = table.rows[row_idx]
    for i, text in enumerate(texts):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, color)


def add_data_row(table, row_idx, texts, bold_first=True):
    row = table.rows[row_idx]
    for i, text in enumerate(texts):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.font.size = Pt(10)
        if i == 0 and bold_first:
            run.bold = True


def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return h


def add_para(text, bold=False, indent=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    if bold:
        run.bold = True
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.7)
    return p


# ============================================================
#                       封 面
# ============================================================
for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('软件工程实训\n小组项目报告')
run.font.size = Pt(28)
run.bold = True
run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('ZUUL — 失落的古迹')
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

doc.add_paragraph()

cover_info = [
    ('项 目 名 称', 'ZUUL — 失落的古迹（Roguelike 地牢探险游戏）'),
    ('课 程 名 称', '软件工程实训'),
    ('指 导 教 师', '唐祖锴'),
    ('小 组 名 称', 'AAA603 小组'),
    ('学    院', '武汉理工大学 · 计算机科学与技术学院'),
]
for label, value in cover_info:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'{label}：{value}')
    run.font.size = Pt(14)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f'报告日期：{datetime.date.today().strftime("%Y年%m月%d日")}')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()

# ============================================================
#                   1. 任务概述
# ============================================================
add_heading('1. 任务概述', level=1)

add_para(
    '本次软件工程实训的任务是开发一款基于 Web 的 Roguelike 地牢探险游戏 —— ZUUL — 失落的古迹。'
    '项目以经典文本冒险游戏 World of Zuul 为基础，进行图形化、实时交互化的全面扩展与重构。',
    indent=True
)

add_para('实践/项目/系统的主要内容：', bold=True)
contents = [
    '设计并实现一个程序化随机地图生成系统，生成 10×15 网格（约30-50个房间）的迷宫地图',
    '构建基于 Phaser 3 游戏引擎的实时图形化前端界面，支持键盘操作、小地图导航、粒子特效',
    '开发基于 Spring Boot 3.2.0 的后端 RESTful API 服务，提供完整的游戏逻辑与数据管理',
    '实现三种攻击模式（扇形扫击、直线突刺、蓄力360°）及空间命中判定算法',
    '设计五种装备槽位系统（武器、护甲、披风、戒指、项链）及属性加成体系',
    '实现四种状态效果系统（烧伤、中毒、流血、天使祝福），支持层数叠加与独立结算',
    '构建商店交易系统、篝火祭坛交互系统、奇遇事件系统、铁匠强化系统等多个子系统',
    '实现基于 H2 数据库的完整游戏存档功能，支持保存、读取、删除等多槽位管理',
]
for c in contents:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(c)
    run.font.size = Pt(11)

add_para('任务目的：', bold=True)
add_para(
    '通过本次实训，巩固软件工程规范，提高面向对象建模与抽象能力，培养小组协同开发能力。'
    '具体包括：掌握前后端分离架构的设计与实现；实践 Git 分支管理与代码审查流程；'
    '学习 RESTful API 设计规范；熟悉 Spring Boot 与 Vue 3 的完整开发流程；'
    '体验真实团队协作中的任务拆分、进度管理、问题追踪等软件工程实践。',
    indent=True
)

doc.add_page_break()

# ============================================================
#               2. 小组任务计划
# ============================================================
add_heading('2. 小组任务计划', level=1)

# 2.1 任务分析
add_heading('2.1 任务分析', level=2)
add_para(
    '本项目作为一个完整的 Roguelike 地牢探险游戏，需要同时覆盖后端逻辑、前端渲染、'
    '数据库持久化等多个技术领域。经过需求分析与技术评估，项目工作重点如下：',
    indent=True
)

add_para('工作重点：', bold=True)
key_points = [
    '核心游戏机制：玩家移动、战斗系统（三种攻击模式 + 空间判定）、怪物 AI（检测/攻击/击退）、物品拾取/丢弃/使用',
    '随机内容生成：程序化地图生成、随机商店商品、随机祭坛恩赐、随机奇遇事件（8种事件类型）',
    '状态与装备系统：烧伤/中毒/流血/迟缓/束缚/天使祝福六种状态效果、五槽位装备系统（属性加成交互计算）',
    '数据持久化：基于 Spring Data JPA + H2 的完整存档系统，保存全部游戏状态到数据库',
    '前端游戏引擎集成：Vue 3 组件化界面与 Phaser 3 游戏引擎的嵌套集成、CustomEvent 通信机制',
    'UI/UX 优化：小地图导航、控制面板（ESC打开/关闭）、背包面板、攻击粒子特效与屏幕振动',
]
for k in key_points:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(k)
    run.font.size = Pt(11)

# 2.2 项目架构与功能设计
add_heading('2.2 项目架构与功能设计', level=2)
add_para(
    '项目采用前后端分离的 B/S 架构。后端基于 Java 17 + Spring Boot 3.2.0 提供 13 个 RESTful API 端点（含风隐/寒冰风暴），'
    '前端基于 Vue 3.3.4 + Phaser 3.60.0 构建图形化游戏界面。前后端通过 HTTP JSON 通信。',
    indent=True
)

add_para('后端架构（Spring Boot + Maven）：', bold=True)
be_items = [
    ('Controller 层', 'GameController.java — 13 个 REST API 端点（含风隐/寒冰风暴）'),
    ('Service 层', 'GameService.java（命令分发/攻击判定/风隐/寒冰风暴/状态驱动）+ SaveService.java（存档序列化）'),
    ('Model 层', 'Room/Player/Monster/Bag/Item/Status/WisdomBoon/Magic/Money 等 18 个数据模型'),
    ('Command 层', 'GoCommand/AttackCommand/BagCommand/ShopCommand/InteractCommand/WaveCommand/MonsterAttackCommand 等 11 个命令'),
    ('Repository 层', '5 个 Spring Data JPA Repository 接口'),
]
table = doc.add_table(rows=len(be_items) + 1, cols=2)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
make_header_row(table, 0, ['层级', '说明'])
for idx, (layer, desc) in enumerate(be_items):
    add_data_row(table, idx + 1, [layer, desc])

doc.add_paragraph()

add_para('前端架构（Vue 3 + Vite + Phaser 3）：', bold=True)
fe_items = [
    ('组件层', 'App.vue（主应用含存档弹窗）、MainMenu.vue（主菜单）、GameCanvas.vue（Phaser游戏画布）、BackpackPanel.vue（背包面板）、ControlPanel.vue（ESC控制面板）、MinimapPanel.vue（小地图）'),
    ('Composables', 'useApi.js（API调用+风隐/寒冰风暴封装）、useBackpack.js（背包逻辑）、useKeyboard.js（键盘控制+B/F/ESC键管理）、useControlPanel.js（控制面板单例）、useSlotDialog.js（存档弹窗）'),
    ('游戏引擎', 'EntityDrawer.js（游戏实体绘制器：人物/怪物/掉落物/祭坛/商人）、game/constants.js（全局常量/攻击/技能/怪物/渲染配置）'),
    ('API 层', 'api/gameApi.js（存档CRUD封装）'),
]
table = doc.add_table(rows=len(fe_items) + 1, cols=2)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
make_header_row(table, 0, ['层级', '说明'])
for idx, (layer, desc) in enumerate(fe_items):
    add_data_row(table, idx + 1, [layer, desc])

doc.add_paragraph()

add_para('核心功能模块：', bold=True)
modules = [
    '随机地图生成（GenerateRoom.java）：10×15 网格，7种房间类型，种子随机数保证存档一致性',
    '实时战斗系统：Sweep（扇形120px/135°）、Pierce（直线21px/120px，无敌帧180ms）、Charged（360°/150px）三种攻击',
    '怪物系统：TYPE_NORMAL(0)、TYPE_ELITE(1)、TYPE_BOSS(2)，特殊 FLAME_SLIME（死后自爆烧伤）/ SLIME（触碰迟缓）/ 骷髅（中毒）/ 狼人（流血）/ 食人魔（吸血）',
    '状态效果系统：BURN(3s魔法伤害)、POISON(1s真实伤害)、BLEED(攻击触发)、SLOW(10s移速减半)、BIND(3s无法移动)、ANGEL_BUFF(30s全属性×1.5)',
    '背包与装备系统：消耗品合并显示、5槽位装备（weapon/armor/cloak/ring/amulet）',
    '商店系统：6件随机商品、半价回收',
    '篝火祭坛：HEAL(50%回复)/TRAIN(25%属性提升)/WISDOM(3选1恩赐)',
    '奇遇事件系统：8种事件（WOODEN_CHEST/GOLDEN_CHEST/ELITE_ENEMY/FOUNTAIN/CHEST/MAIDEN/ANGEL/BLACKSMITH）',
    '铁匠强化系统：装备属性提升至150%',
    '月光波技能：消耗20MP，单体魔法攻击，弹射2次，击退效果',
    '风隐形态：移速+100%、免疫负面状态，每秒消耗2MP',
    '寒冰风暴：消耗25MP，全房间AOE 3次100%法伤+10秒迟缓',
    '存档系统：H2数据库持久化，完整游戏状态保存/读取/删除',
]
for m in modules:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(m)
    run.font.size = Pt(11)

# 2.3 团队组成与分工
add_heading('2.3 团队组成与分工', level=2)
add_para('AAA603 小组由 3 名成员组成，分工如下：', indent=True)

member_table = doc.add_table(rows=4, cols=4)
member_table.style = 'Table Grid'
member_table.alignment = WD_TABLE_ALIGNMENT.CENTER
make_header_row(member_table, 0, ['姓名', 'GitHub', '分支', '主要职责'])
member_data = [
    ('李冬晨（组长）',
     '@Sader-Lee',
     'ldc',
     '核心后端开发：项目架构设计、前后端分离、怪物系统（生成/攻击/掉落）、商店系统（购买/出售）、背包系统、房间战斗封锁、数据库持久化（H2+JPA）、前后端联调'),
    ('蒋志成',
     '@BytesJiang',
     'jzc',
     '前端开发：玩家属性系统设计与开发、小地图模块开发、UI/UX 优化与视觉提升、玩家攻击系统设计与实现、玩家技能系统设计与实现、状态效果系统设计与实现、随机地图生成系统开发、多类型房间功能开发、前端架构重构与优化、缺陷修复与性能调优'),
    ('熊俊森',
     '@XJS-123',
     'xjs',
     '特殊系统开发：主菜单页面搭建设计、饰品系统完整实现、人物模型绘制、游戏实体构造（掉落物实体、EntityDrawer）、地图背景添加、项目文档生成（API.md/README.md/REPORT.docx）'),
]
for idx, (name, github, branch, duty) in enumerate(member_data):
    add_data_row(member_table, idx + 1, [name, github, branch, duty])

# 2.4 软件工具与项目约束
add_heading('2.4 软件工具与项目约束', level=2)

add_para('开发工具：', bold=True)
tools = [
    ('IntelliJ IDEA', '后端 Java 开发'),
    ('VS Code', '前端 Vue/Phaser 开发'),
    ('GitHub', '代码仓库与协同开发'),
    ('GitHub Issues', '任务管理与分配'),
    ('Postman', 'API 测试与调试'),
]
table = doc.add_table(rows=len(tools) + 1, cols=2)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
make_header_row(table, 0, ['工具', '用途'])
for idx, (tool, usage) in enumerate(tools):
    add_data_row(table, idx + 1, [tool, usage])

doc.add_paragraph()
add_para('项目管理策略与强制约束：', bold=True)
constraints = [
    '版本控制：必须使用 Git 进行版本管理，所有代码提交需有清晰的信息描述',
    '分支策略：采用 master/ldc/jzc/xjs 多分支开发，通过 Pull Request 合并代码',
    '代码审查：每个 PR 需至少一名其他成员审查通过后才能合并',
    '代码风格：后端遵循 Java 标准命名规范，前端遵循 ESLint 规则',
    'API 设计：所有接口遵循统一响应格式 {status, message, data}',
    '前后端通信：所有游戏状态由后端维护，前端只做呈现与用户输入捕获',
    '文档同步：代码更新后需同步更新 API.md 和 README.md',
]
for c in constraints:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(c)
    run.font.size = Pt(11)

doc.add_page_break()

# ============================================================
#               3. 小组软件过程
# ============================================================
add_heading('3. 小组软件过程', level=1)

# 3.1 代码仓库
add_heading('3.1 代码仓库', level=2)
add_para(
    '项目代码托管于 GitHub 私有仓库，地址为：https://github.com/WHUT-ZUUL/kai-fa-aaa603。'
    '仓库采用统一组织管理，所有成员均有推送权限。仓库根目录包含后端（backend/）、前端（frontend/）、'
    '开发计划（plans/）、项目文档（README.md、API.md、REPORT.docx）等核心目录。',
    indent=True
)

# 3.2 软件版本计划
add_heading('3.2 软件版本计划', level=2)
add_para('项目采用迭代式开发，整体分为以下阶段：', indent=True)

version_plan = [
    ('v0.1 — 基础框架', '搭建 Spring Boot + Vue 3 项目骨架，实现基本的前后端通信，完成玩家移动和房间探索功能'),
    ('v0.2 — 战斗系统', '实现实时战斗系统（三种攻击模式）、怪物 AI 和怪物生成机制'),
    ('v0.3 — 背包与装备', '实现背包系统（拾取/使用/丢弃）、五槽位装备系统（装备/卸下/属性加成）'),
    ('v0.4 — 商店与交互', '实现商店系统（购买/出售）、篝火祭坛系统（治愈/训练/博学）、奇遇事件系统'),
    ('v0.5 — 状态效果', '实现烧伤/中毒/流血/天使祝福四种状态效果系统，支持层数叠加'),
    ('v0.6 — 存档与 UI', '实现基于 H2 数据库的完整存档系统、控制面板(ESC)、小地图组件、UI/UX 优化'),
    ('v1.0 — 集成与交付', '前后端集成测试、Bug 修复、代码优化与冗余清理、项目文档生成'),
]
for ver, desc in version_plan:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(f'{ver}：{desc}')
    run.font.size = Pt(11)

# 3.3 开发分支模型与提交流程
add_heading('3.3 开发分支模型与提交流程', level=2)
add_para(
    '项目采用三分支模型：master（主分支）、ldc（后端核心）、jzc（前端开发）、xjs（特殊系统与实体）。'
    '所有新功能的开发在各自独立分支上进行，完成后发起 Pull Request 到 master 分支。',
    indent=True
)

branch_flow = [
    '功能开发：各成员在个人分支（ldc/jzc/xjs）上开发新功能，每日多次提交',
    '代码同步：定期从 master 合并最新代码到个人分支，保持代码同步',
    'Pull Request：功能完成后，在 GitHub 上发起 PR 到 master 分支',
    '代码审查：至少一名其他成员审查代码，提出改进意见',
    '合并：审查通过后，由 PR 创建者或组长合并到 master',
    '标签管理：合并后在 master 分支打上版本标签（如 v0.6）',
]
for bf in branch_flow:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(bf)
    run.font.size = Pt(11)

# 3.4 代码规范与检查流程
add_heading('3.4 代码规范与检查流程', level=2)
add_para(
    '后端 Java 代码遵循阿里巴巴 Java 开发手册规范，前端代码遵循 Vue 3 Composition API 风格指南。'
    '代码审查重点检查以下方面：',
    indent=True
)
code_review = [
    '命名规范：类名 PascalCase、方法名 camelCase、常量 UPPER_SNAKE_CASE',
    '代码整洁：无冗余导入、无用代码注释、控制台输出',
    '异常处理：所有可能的空指针需做 null 检查，错误信息有意义',
    'API 一致性：响应格式必须遵循统一的 {status, message, data} 结构',
    '安全性：存档操作需校验 saveId 合法性',
]
for cr in code_review:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(cr)
    run.font.size = Pt(11)

# 3.5 项目测试计划与实施流程
add_heading('3.5 项目测试计划与实施流程', level=2)
add_para(
    '项目采用"开发自测 + 集成联调"的测试策略。由于项目周期较短，未引入自动化测试框架，'
    '主要通过以下方式进行测试：',
    indent=True
)
test_items = [
    '后端 API 测试：使用 Postman 对各 API 端点进行请求/响应验证',
    '前端功能测试：在浏览器中手动测试各交互功能（移动/攻击/拾取/商店/存档等）',
    '前后端联调：在开发过程中持续进行前后端联调，确保 JSON 数据格式一致',
    '边界测试：测试空背包使用物品、在非商店房间使用 shop 命令、金钱不足时购买等边界情况',
    '存档兼容性测试：验证存档/读档后游戏状态的一致性',
]
for t in test_items:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(t)
    run.font.size = Pt(11)

# 3.6 项目集成与发布规范
add_heading('3.6 项目集成与发布规范', level=2)
add_para(
    '项目通过 GitHub 进行持续集成与发布管理。各成员在个人分支完成开发和自测后，通过 Pull Request '
    '将代码合并到 master 主分支。每次合并前需确保代码可编译、可运行，且不破坏现有功能。',
    indent=True
)
release_items = [
    '集成触发：每次 Pull Request 合并到 master 后触发集成验证',
    '构建验证：后端使用 Maven 编译验证（mvn compile），前端使用 Vite 构建验证（npm run build）',
    '发布流程：代码在 master 分支上经过验证后，打上版本标签进行发布',
    '文档同步：版本发布时同步更新 README.md 和 API.md 中的版本号与更新日期',
]
for r in release_items:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(r)
    run.font.size = Pt(11)

# 3.7 其他过程说明
add_heading('3.7 其他过程说明', level=2)
add_para(
    '开发过程中使用了 GitHub Issues 进行任务管理和分配。每个功能点对应一个 Issue，分配给对应成员后'
    '在个人分支上实现。Issues 支持标签分类（enhancement/bug/documentation）和里程碑管理。'
    '项目前期以功能开发为主（约85%的 Issues），后期转向 Bug 修复和代码优化（约15%的 Issues）。',
    indent=True
)

doc.add_page_break()

# ============================================================
#               4. 小组成员任务报告
# ============================================================
add_heading('4. 小组成员任务报告', level=1)

# 4.1 成员一：李冬晨
add_heading('4.1 成员一（李冬晨）任务分工及实施过程描述', level=2)
add_para('角色：组长 / 核心后端开发', bold=True)
add_para('GitHub：@Sader-Lee | 开发分支：ldc', bold=True)

add_para('任务分工：', bold=True)
ldc_tasks = [
    '项目架构设计与技术选型：确定 Spring Boot + Vue 3 + Phaser 3 的技术栈，实现前后端分离架构',
    '核心游戏模型开发：实现 Player、Monster、Room 等 18 个数据模型',
    '怪物系统：怪物生成（spawnMonsters 三种类型）、怪物攻击、怪物掉落（processMonsterDrop）',
    '商店系统：ShopCommand 完整实现（buy/sell），商品管理、货币扣除、折价回收（买价一半）',
    '背包系统：Bag 类的 add/remove/use/discard 操作，消耗品合并与装备独立展示',
    '房间战斗封锁：怪物房间 exit 封锁/解锁机制',
    '数据库持久化：基于 Spring Data JPA 的完整存档系统，5 个 Repository 接口',
    '前后端联调与攻击特效',
]
for t in ldc_tasks:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(t)
    run.font.size = Pt(11)

add_para('实施成果：完成了后端核心系统的所有功能，包括 13 个 API 端点、11 个命令处理器、'
         '18 个数据模型、5 个 Repository 接口。实现了游戏的核心逻辑闭环。', indent=True)

# 4.2 成员二：蒋志成
add_heading('4.2 成员二（蒋志成）任务分工及实施过程描述', level=2)
add_para('角色：前端开发 / 战斗系统', bold=True)
add_para('GitHub：@BytesJiang | 开发分支：jzc', bold=True)

add_para('任务分工：', bold=True)
jzc_tasks = [
    '玩家属性系统设计与开发：实现玩家属性管理、装备加成计算、伤害计算等核心机制',
    '小地图模块开发：实现 MinimapPanel.vue 组件，支持按房间类型颜色区分',
    'UI/UX 优化与视觉提升：攻击特效（火焰粒子、屏幕振动、蓄力特效、残影渐隐）、Buff/Debuff 图标显示',
    '玩家攻击系统设计与实现：Sweep（扇形/135°）、Pierce（直线/突刺/击退）、Charged（蓄力/360°）三种攻击模式',
    '玩家技能系统设计与实现：月光波技能（弹射 2 次、击退）、风隐形态、寒冰风暴',
    '状态效果系统设计与实现：烧伤/中毒/流血/迟缓/束缚的前端图标显示、倒计时、闪烁效果',
    '随机地图生成系统开发：实现 10×15 网络的地图生成逻辑，7种房间类型分布',
    '多类型房间功能开发：篝火祭坛交互展示、奇遇房间界面、商店界面',
    '前端架构重构与优化：常量提取到 constants.js、API 封装到 gameApi.js、组件拆分（BackpackPanel.vue / ControlPanel.vue / MinimapPanel.vue）',
    '缺陷修复与性能调优',
]
for t in jzc_tasks:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(t)
    run.font.size = Pt(11)

add_para('实施成果：构建了完整的 Vue 3 + Phaser 3 前端游戏界面，包括 5 个核心组件、5 个 composables、'
         '全局常量系统、攻击特效系统、小地图导航系统。完成了前后端分离的通信架构。', indent=True)

# 4.3 成员三：熊俊森
add_heading('4.3 成员三（熊俊森）任务分工及实施过程描述', level=2)
add_para('角色：特殊系统开发 / 实体与文档', bold=True)
add_para('GitHub：@XJS-123 | 开发分支：xjs', bold=True)

add_para('任务分工：', bold=True)
xjs_tasks = [
    '主菜单页面搭建设计：MainMenu.vue 界面实现',
    '饰品系统完整实现：5 槽位装备系统（weapon/armor/cloak/ring/amulet），装备/卸下逻辑，属性加成交互计算',
    '人物模型绘制：游戏角色视觉呈现',
    '游戏实体构造：EntityDrawer.js 实现掉落物可视化、DroppedItem 坐标管理、祭坛/商人绘制',
    '地图背景添加：4种背景图（bg0-bg3）对应不同房间类型',
    '项目文档生成：API.md、README.md、REPORT.docx 的编写与更新',
]
for t in xjs_tasks:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(t)
    run.font.size = Pt(11)

add_para('实施成果：完成了主菜单页面、饰品系统、人物模型绘制、地图背景添加等前端模块，'
         '生成了完整的项目文档。', indent=True)

# 4.4 成员四
add_heading('4.4 成员四任务分工及实施过程描述', level=2)
add_para('（本小组共 3 名成员，此章节不适用）', indent=True)

doc.add_page_break()

# ============================================================
#           5. 实施过程问题记录与分析
# ============================================================
add_heading('5. 实施过程问题记录与分析', level=1)

add_para('在项目合作开发过程中，遇到了以下主要问题及其分析与解决方案：', indent=True)

problems = [
    ('前后端数据同步问题',
     '问题描述：攻击判定涉及前端坐标与后端模型的同步。前端 Phaser 引擎使用像素坐标，'
     '后端需要根据坐标做空间命中判定，双方坐标系统不一致会导致判定偏差。',
     '分析：攻击判定需要在服务端统一执行以保证公平性和存档一致性。前端发送坐标快照到后端，'
     '后端执行空间几何运算（扇形角度判定/直线距离判定/圆形范围判定）。',
     '解决方案：前端在攻击时将玩家坐标、怪物坐标快照发送到后端，后端 GameService.performAttack '
     '方法统一执行命中判定。实现了三种攻击算法的空间命中检测（isHit 方法）。'),
    ('状态效果系统设计',
     '问题描述：多种 Buff/Debuff 需要独立结算但又可能同时存在，且支持层数叠加。'
     '前端需要实时显示剩余时间、层数等信息。',
     '分析：需要统一的 StatusManager 管理中心，每种状态效果有独立的 tick 逻辑和层数衰减规则。',
     '解决方案：Status.StatusManager 统一管理所有效果，采用层数叠加机制。BURN 每3秒结算（层数减半），'
     'POISON 每秒结算（层数减1），BLEED 攻击触发（层数减1），ANGEL_BUFF 持续30秒自动结束。'
     '前端通过 getActiveEffectsInfo() 获取当前活跃效果列表。'),
    ('随机种子一致性问题',
     '问题描述：存档读档时需要保证随机地图的可复现性，否则读档后地图与存档时不符。',
     '分析：Java 的 Random 类使用种子值初始化，相同种子产生相同的随机序列。'
     '需要在存档时保存种子值，读档时使用同样的种子重建地图和怪物。',
     '解决方案：Game 类使用固定种子初始化（new Random(seed)），存档时保存种子到数据库，'
     '读档时恢复种子重建完整的房间网络和怪物分布。'),
    ('装备同槽位冲突处理',
     '问题描述：装备新武器时，旧武器需要自动卸下；卸下装备后属性要恢复。'
     '不同物品对应不同槽位（铁剑→weapon、暗影披风→cloak 等）。',
     '分析：需要建立物品名称到槽位的映射表，equip 时自动检测同槽位旧装备。',
     '解决方案：Player.getItemSlot() 方法根据物品名称中文关键词判断槽位（剑→weapon、盾→armor、'
     '披风→cloak、戒指→ring、项链→amulet）。equipItem 方法先检查同槽位，调用 unequipSlot 卸下旧装备，'
     '再应用新装备属性。'),
    ('中文字符编码与前端显示',
     '问题描述：后端响应中的中文字符在控制台打印时出现编码错误（GBK 编码问题）。',
     '分析：Windows 系统控制台默认 GBK 编码，Python 脚本输出的 emoji 和特殊字符无法编码。',
     '解决方案：脚本输出使用 ASCII 安全字符替代 emoji，后端 JSON 默认 UTF-8 不受影响。'),
    ('存档写权限冲突',
     '问题描述：IDEA 开发工具锁定了 REPORT.docx 文件，导致 Python 脚本无法覆盖写入。',
     '分析：IDEA 打开了文件导致文件被进程占用。',
     '解决方案：关闭 IDEA 中的文件标签后再执行脚本，或先写入临时文件再替换。'),
]
for title, desc, analysis, solution in problems:
    add_para(title, bold=True)
    add_para(f'问题描述：{desc}', indent=True)
    add_para(f'分析：{analysis}', indent=True)
    add_para(f'解决方案：{solution}', indent=True)
    doc.add_paragraph()

doc.add_page_break()

# ============================================================
#               6. 任务总结
# ============================================================
add_heading('6. 任务总结', level=1)

add_para('6.1 项目完成情况', bold=True)
add_para(
    '本项目已按计划完成全部功能开发，通过了多次前后端联调测试，项目状态稳定。'
    '已实现的功能包括：程序化随机地图生成、实时战斗系统（三种攻击模式）、怪物系统（三种类型+火焰史莱姆）、'
    '状态效果系统（四种效果）、背包与五槽位装备系统、商店交易系统、篝火祭坛系统（三种祭坛+八种恩赐）、'
    '奇遇事件系统（八种事件）、铁匠强化系统、月光波技能、完整存档系统（H2数据库持久化）。'
    '前端完成了 Vue 3 组件化架构（6个组件/5个composables）、Phaser 3 游戏引擎集成、'
    '小地图导航、控制面板、攻击粒子特效、屏幕振动等 UI/UX 功能。',
    indent=True
)

add_para('6.2 经验与收获', bold=True)
gains = [
    '深入理解了前后端分离架构的设计思想与实践方法，掌握了 Spring Boot + Vue 3 的完整开发流程',
    '掌握了 Phaser 3 游戏引擎的基本用法：场景管理、碰撞检测、粒子系统、动画与特效',
    '学习了 Roguelike 游戏的程序化内容生成技术：随机地图生成、随机事件系统、平衡性设计',
    '实践了 RESTful API 设计规范：统一响应格式、错误处理、状态管理',
    '掌握了 Spring Data JPA 的实体关系映射与复杂状态序列化/反序列化',
    '通过 Git 分支管理与代码审查，体验了真实团队协作的开发流程',
    '培养了问题分析与解决能力：从坐标同步到状态管理、从种子一致性到装备冲突处理',
]
for g in gains:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(g)
    run.font.size = Pt(11)

add_para('6.3 项目亮点', bold=True)
highlights = [
    '完整的空间命中判定系统：后端实现三种攻击模式的几何命中算法（扇形/直线/全向），前端发送坐标快照做服务端校验',
    '丰富的状态效果系统：四种效果均支持层数叠加与独立结算，前端实时展示状态图标与倒计时',
    '灵活的装备系统：五槽位设计支持武器/护甲/披风/戒指/项链的装备与卸下，属性加成实时反映',
    '多维度的内容生成：随机地图 + 随机商店 + 随机祭坛恩赐 + 随机奇遇事件，保证每次游戏的全新体验',
    '完整的存档系统：基于 H2 数据库实现全量游戏状态持久化，支持多槽位存档管理',
    '前端组件化重构：将背包面板、控制面板、小地图拆分为独立 Vue 组件，通过 CustomEvent 通信',
]
for h in highlights:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(h)
    run.font.size = Pt(11)

add_para('6.4 未来展望', bold=True)
future = [
    '增加更多怪物种类与 Boss 技能机制',
    '引入装备强化/镶嵌/附魔系统',
    '增加游戏内成就系统与排行榜',
    '优化 Phaser 性能，支持更多粒子效果与动画帧',
    '添加音效与背景音乐系统',
    '支持多人在线联机模式',
]
for f in future:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(f)
    run.font.size = Pt(11)

# ============================================================
#                       结 尾
# ============================================================
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('— 报告完 —')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f'武汉理工大学 · 计算机科学与技术学院\n{datetime.date.today().strftime("%Y年%m月%d日")}')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

# ========== 保存 ==========
output_path = 'REPORT.docx'
doc.save(output_path)
print(f'[OK] Report generated: {output_path}')
